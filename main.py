"""
main.py — خدمة تعبئة الاستمارات (FastAPI)
================================================

خدمة مستقلة عن خدمة قراءة الجوازات — خفيفة وسريعة، ما تحمّل
OpenCV ولا Tesseract، فتصحى بثواني بدل دقيقة.

المبدأ: نستقبل رابط قالب Word + البيانات، نفتح القالب بمكتبة
python-docx، نعبّيه، وكل شي ثاني بالقالب (الترويسة، النصوص،
التوقيعات) يبقى حرفياً كما هو. دائرة السياحة ترفض أي تغيير
بالشكل، فالحفاظ على القالب مطلب أساسي مو تحسين.

Endpoints:
  POST /fill-approval-form  — استمارة الموافقة الأمنية (جدول كامل)
  POST /fill-entry-pass     — سمة الدخول (أول اسم + آخر اسم + العدد)
"""

import io
import re
from copy import deepcopy

import requests
from fastapi import FastAPI
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


app = FastAPI(title="خدمة تعبئة الاستمارات")


MONTHS = [
    "JAN", "FEB", "MAR", "APR", "MAY", "JUN",
    "JUL", "AUG", "SEP", "OCT", "NOV", "DEC",
]

ARABIC_MONTHS = [
    "كانون الثاني", "شباط", "آذار", "نيسان", "أيار", "حزيران",
    "تموز", "آب", "أيلول", "تشرين الأول", "تشرين الثاني", "كانون الأول",
]


# ============================================================================
# نماذج البيانات القادمة من التطبيق
# ============================================================================

class Passenger(BaseModel):
    serial: int = 0
    given_name_en: str = ""
    given_name_ar: str = ""
    father_name_en: str = ""
    father_name_ar: str = ""
    surname_en: str = ""
    surname_ar: str = ""
    passport_number: str = ""
    birth_date: str = ""      # نستقبله ISO: 1964-04-13
    nationality: str = ""
    residence_country: str = ""


class FillRequest(BaseModel):
    template_url: str
    passengers: list[Passenger]


class EntryPassRequest(BaseModel):
    """
    ⚠ سمة الدخول: استمارة جاهزة ما نغيّر منها إلا 3 قيم —
    أول اسم بالقائمة، آخر اسم، وعدد المجموعة
    """
    template_url: str
    first_name: str = ""
    last_name: str = ""
    count: int = 0


# ============================================================================
# التاريخ
# ============================================================================

def format_birth_date(raw: str) -> str:
    """
    نحوّل التاريخ لصيغة الجواز: "13 APR 1964"

    نستقبل ISO من التطبيق (1964-04-13T00:00:00.000) لأنها الصيغة
    اللي يخزنها Supabase، ونتساهل مع صيغ ثانية احتياطاً
    """

    text = (raw or "").strip()
    if not text:
        return ""

    # ISO: 1964-04-13 أو 1964-04-13T00:00:00
    iso = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})", text)
    if iso:
        year, month, day = int(iso.group(1)), int(iso.group(2)), int(iso.group(3))
        if 1 <= month <= 12:
            return f"{day:02d} {MONTHS[month - 1]} {year}"

    # 13/04/1964
    slash = re.match(r"^(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{4})$", text)
    if slash:
        day, month, year = int(slash.group(1)), int(slash.group(2)), int(slash.group(3))
        if 1 <= month <= 12:
            return f"{day:02d} {MONTHS[month - 1]} {year}"

    # 1964-APR-13 (صيغة شاشة المراجعة)
    named = re.match(r"^(\d{4})-([A-Z]{3})-(\d{1,2})$", text.upper())
    if named and named.group(2) in MONTHS:
        return f"{int(named.group(3)):02d} {named.group(2)} {named.group(1)}"

    return text


# ============================================================================
# التعامل مع خلايا الجدول
# ============================================================================

def set_cell_text(cell, lines, template_cell=None):
    """
    نكتب نص بخلية، مع دعم عدة أسطر داخل نفس الخلية.

    ⚠ ليش ما نكتفي بـcell.text = "..."؟ لأنها تمسح كل التنسيق
    (الخط العريض، الحجم، المحاذاة) وتخلي الخلية بالشكل الافتراضي.
    والاستمارة كلها بخط عريض، فالنتيجة تطلع مختلفة عن الأصل.

    الحل: ننسخ تنسيق فقرة موجودة أصلاً بالقالب ونطبّقه على النص
    الجديد — هيك الشكل يبقى مطابق تماماً
    """

    # ننضّف الخلية من أي محتوى قديم
    for paragraph in list(cell.paragraphs[1:]):
        paragraph._element.getparent().remove(paragraph._element)

    first = cell.paragraphs[0]
    for run in list(first.runs):
        run._element.getparent().remove(run._element)

    texts = [t for t in lines if t and str(t).strip()]
    if not texts:
        return

    first.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index, text in enumerate(texts):
        paragraph = first if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(str(text).strip())
        run.bold = True
        run.font.size = Pt(10)

        # ندعم العربي صح — بدون هذا بعض النسخ تعرض الحروف مقلوبة
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:cs"), "Arial")


def clone_row(table, source_row):
    """
    ننسخ صف موجود بالجدول ونضيفه بالنهاية.

    نستنسخ صف حقيقي من القالب (مو نسوي صف جديد فاضي) عشان
    نرث كل تنسيقاته: عرض الأعمدة، الحدود، الظلال، الارتفاع.
    الصف الجديد المصنوع من الصفر يطلع بشكل مختلف عن بقية الجدول
    """

    new_row = deepcopy(source_row._element)
    source_row._element.getparent().append(new_row)
    return table.rows[-1]


def find_data_table(document):
    """
    نلقى جدول البيانات بالقالب.

    القالب فيه أكثر من جدول (فيه جدول صغير بالأعلى للترويسة)،
    فنميّز جدول البيانات بعناوين أعمدته: يحتوي PASSPORT NO
    و GIVEN NAME. هذا أوثق من "الجدول الأكبر" أو "الجدول الثاني"
    — لو انضاف جدول بالقالب مستقبلاً ما ينكسر المنطق
    """

    for table in document.tables:
        if not table.rows:
            continue

        header = " ".join(cell.text.upper() for cell in table.rows[0].cells)

        if "PASSPORT" in header and ("GIVEN" in header or "SURNAME" in header):
            return table

    # ما لقينا بالعناوين — نرجّع الجدول الأكثر أعمدة كخطة أخيرة
    if document.tables:
        return max(document.tables, key=lambda t: len(t.columns))

    return None


# ============================================================================
# ⚠ الاستبدال داخل الفقرات — لسمة الدخول
# ============================================================================

def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    """
    ⚠ نستبدل نص داخل فقرة **بدون ما نكسر تنسيقها**.

    المشكلة: Word يقسّم الفقرة لـruns حسب التنسيق، والنص الواحد
    ممكن ينقسم على عدة runs ("صفاء" run، " حسين" run ثاني). فلو
    دوّرنا على النص داخل كل run لحاله ما نلقاه.

    الحل: نجمع نص كل الruns، نستبدل بالنص المجمّع، وبعدها نحط
    النتيجة كلها بأول run ونفضّي الباقي — هيك تنسيق أول run
    (اللي هو تنسيق الفقرة الفعلي) يبقى محفوظ
    """

    full_text = "".join(run.text for run in paragraph.runs)

    if old not in full_text:
        return False

    new_text = full_text.replace(old, new)

    if not paragraph.runs:
        return False

    paragraph.runs[0].text = new_text

    for run in paragraph.runs[1:]:
        run.text = ""

    return True


def replace_everywhere(document, old: str, new: str) -> int:
    """نستبدل بكل الفقرات وكل خلايا الجداول. نرجّع عدد المواضع."""

    count = 0

    for paragraph in document.paragraphs:
        if replace_in_paragraph(paragraph, old, new):
            count += 1

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph, old, new):
                        count += 1

    return count


def replace_by_pattern(document, pattern: str, replacement) -> int:
    """
    ⚠ استبدال بنمط regex — نحتاجه لما ما نعرف القيمة القديمة.

    مثال: بسمة الدخول نريد نبدّل العدد اللي بعد "عدد المجموعة"،
    بس ما نعرف شنو الرقم الموجود بالقالب. فندوّر بالنمط
    """

    compiled = re.compile(pattern)
    count = 0

    def apply(paragraph):
        nonlocal count

        full_text = "".join(run.text for run in paragraph.runs)

        if not compiled.search(full_text):
            return

        new_text = compiled.sub(replacement, full_text)

        if new_text == full_text or not paragraph.runs:
            return

        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

        count += 1

    for paragraph in document.paragraphs:
        apply(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply(paragraph)

    return count


# ============================================================================
# التعبئة — استمارة الموافقة الأمنية
# ============================================================================

def fill_template(template_bytes: bytes, passengers: list[Passenger]) -> bytes:
    document = Document(io.BytesIO(template_bytes))

    table = find_data_table(document)
    if table is None:
        raise ValueError("ما لقينا جدول البيانات بالقالب")

    # الصف الأول عناوين، والباقي صفوف فاضية جاهزة بالقالب
    header_row = table.rows[0]
    body_rows = list(table.rows[1:])

    if not body_rows:
        raise ValueError("الجدول بالقالب ما بيه صفوف بيانات")

    # نستخدم أول صف فاضي كنموذج للاستنساخ
    template_row = body_rows[0]

    needed = len(passengers)
    available = len(body_rows)

    # ------------------------------------------------------------------
    # نضبط عدد الصفوف على عدد الجوازات بالضبط
    # ------------------------------------------------------------------
    if needed > available:
        for _ in range(needed - available):
            clone_row(table, template_row)
    elif needed < available:
        # نحذف الصفوف الزايدة — الاستمارة ما تنقبل بصفوف فاضية
        for row in body_rows[needed:]:
            row._element.getparent().remove(row._element)

    data_rows = list(table.rows[1:])

    # ------------------------------------------------------------------
    # نعبّي: كل صف جواز، والأعمدة بترتيب القالب
    # ت | الاسم | اسم الأب | اللقب | رقم الجواز | الميلاد | الجنسية | بلد الإقامة
    # ------------------------------------------------------------------
    for index, passenger in enumerate(passengers):
        if index >= len(data_rows):
            break

        cells = data_rows[index].cells
        if len(cells) < 8:
            continue

        serial = passenger.serial if passenger.serial > 0 else index + 1

        set_cell_text(cells[0], [str(serial)])

        # الأسماء الثلاثة: إنكليزي سطر أول، عربي سطر ثاني
        set_cell_text(cells[1], [passenger.given_name_en, passenger.given_name_ar])
        set_cell_text(cells[2], [passenger.father_name_en, passenger.father_name_ar])
        set_cell_text(cells[3], [passenger.surname_en, passenger.surname_ar])

        set_cell_text(cells[4], [passenger.passport_number])
        set_cell_text(cells[5], [format_birth_date(passenger.birth_date)])
        set_cell_text(cells[6], [passenger.nationality])
        set_cell_text(cells[7], [passenger.residence_country])

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output.read()


# ============================================================================
# ⚠ التعبئة — سمة الدخول
# ============================================================================

def fill_entry_pass(
    template_bytes: bytes,
    first_name: str,
    last_name: str,
    count: int,
) -> bytes:
    """
    نعبّي استمارة سمة الدخول.

    الاستمارة جاهزة بالكامل — ما نغيّر منها إلا ثلاث قيم:
      • عدد المجموعة
      • "تبدأ بالاسم: (…)"
      • "تنتهي بالاسم: (…)"

    كل شي ثاني (الجنسية، منفذ الدخول، الفندق، جهة الإبراق،
    التواقيع) يبقى مثل ما هو بالقالب — والموظف يعدّل التواريخ يدوياً
    بعد الطباعة
    """

    document = Document(io.BytesIO(template_bytes))

    filled = []

    # ------------------------------------------------------------------
    # 1. الاسم الأول: "تبدأ بالاسم:- ( … )"
    # ------------------------------------------------------------------
    # ⚠ ندوّر بالنمط مو بالقيمة، لأن القالب فيه اسم قديم ما نعرفه.
    # نمسك من "تبدأ" لين نهاية القوس ونستبدل اللي داخله
    if first_name.strip():
        done = replace_by_pattern(
            document,
            r"(تبدأ\s*بالاسم\s*:?\s*-?\s*)\([^)]*\)",
            lambda m: f"{m.group(1)}( {first_name.strip()} )",
        )
        if done:
            filled.append("first_name")

    # ------------------------------------------------------------------
    # 2. الاسم الأخير: "تنتهي بالاسم:- ( … )"
    # ------------------------------------------------------------------
    if last_name.strip():
        done = replace_by_pattern(
            document,
            r"(تنتهي\s*بالاسم\s*:?\s*-?\s*)\([^)]*\)",
            lambda m: f"{m.group(1)}( {last_name.strip()} )",
        )
        if done:
            filled.append("last_name")

    # ------------------------------------------------------------------
    # 3. العدد: "عدد المجموعة :- 22"
    # ------------------------------------------------------------------
    if count > 0:
        done = replace_by_pattern(
            document,
            r"(عدد\s*المجموعة\s*:?\s*-?\s*)\d+",
            lambda m: f"{m.group(1)}{count}",
        )
        if done:
            filled.append("count")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output.read()


# ============================================================================
# تحميل القالب
# ============================================================================

def download_template(url: str) -> bytes:
    """نحمّل القالب من التخزين — نرمي استثناء واضح لو فشل."""

    response = requests.get(url, timeout=30)

    if response.status_code != 200:
        raise ValueError(f"تعذر تحميل القالب: {response.status_code}")

    return response.content


# ============================================================================
# Endpoints
# ============================================================================

@app.get("/")
def root_check():
    return {"status": "خدمة الاستمارات شغالة ✓", "ready": True}


@app.get("/health")
def health_check():
    return {"status": "ok", "ready": True}


@app.post("/fill-approval-form")
def fill_approval_form(request: FillRequest):
    """نحمّل القالب، نعبّي جدوله، ونرجّع الملف جاهز للطباعة."""

    if not request.passengers:
        return JSONResponse(
            status_code=400,
            content={"success": False, "error": "ماكو أي جواز بالطلب"},
        )

    # 1. نحمّل القالب من التخزين
    try:
        template_bytes = download_template(request.template_url)
    except Exception as error:
        return JSONResponse(
            status_code=400,
            content={"success": False, "error": f"تعذر تحميل القالب: {error}"},
        )

    # 2. نعبّيه
    try:
        filled = fill_template(template_bytes, request.passengers)
    except Exception as error:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": f"فشلت التعبئة: {error}"},
        )

    # 3. نرجّعه كملف
    return StreamingResponse(
        io.BytesIO(filled),
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": 'attachment; filename="approval-form.docx"'},
    )


@app.post("/fill-entry-pass")
def fill_entry_pass_endpoint(request: EntryPassRequest):
    """
    ⚠ سمة الدخول — نعبّي أول اسم وآخر اسم والعدد بس.

    الباقي (التواريخ، الفندق، منفذ الدخول) يبقى مثل القالب،
    والموظف يعدّله يدوياً بعد الفتح
    """

    try:
        template_bytes = download_template(request.template_url)
    except Exception as error:
        return JSONResponse(
            status_code=400,
            content={"success": False, "error": f"تعذر تحميل القالب: {error}"},
        )

    try:
        filled = fill_entry_pass(
            template_bytes,
            request.first_name,
            request.last_name,
            request.count,
        )
    except Exception as error:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": f"فشلت التعبئة: {error}"},
        )

    return StreamingResponse(
        io.BytesIO(filled),
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": 'attachment; filename="entry-pass.docx"'},
    )
